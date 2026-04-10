"""DOCX document renderer using python-docx.

Generates structured resume and cover letter documents in Microsoft
Word format. All blocking I/O from python-docx is offloaded to a
thread-pool executor so the async event loop is never blocked.
"""

from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Any

import structlog

from app.core.exceptions import GenerationError

logger = structlog.get_logger(__name__)


class DOCXRenderer:
    """Renders resume data to DOCX format using python-docx.

    Produces a clean, ATS-friendly Word document with standard
    section formatting (heading + body).
    """

    async def render(
        self,
        template_name: str,
        context: dict[str, Any],
        output_path: Path,
    ) -> Path:
        """Render resume data to a DOCX file.

        Args:
            template_name: Template style name (used for logging).
            context: Resume data dict with keys like ``name``,
                ``email``, ``experience``, ``education``, ``skills``.
            output_path: Where to save the generated DOCX.

        Returns:
            The ``output_path`` on success.

        Raises:
            GenerationError: If document generation fails.
        """
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(
            None,
            self._render_sync,
            template_name,
            context,
            output_path,
        )

    async def render_cover_letter(
        self,
        content: str,
        output_path: Path,
    ) -> Path:
        """Render cover letter text to a DOCX file.

        Args:
            content: Plain-text cover letter content (paragraphs
                separated by double newlines).
            output_path: Where to save the generated DOCX.

        Returns:
            The ``output_path`` on success.

        Raises:
            GenerationError: If document generation fails.
        """
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(
            None,
            self._render_cover_letter_sync,
            content,
            output_path,
        )

    def _render_sync(
        self,
        template_name: str,
        context: dict[str, Any],
        output_path: Path,
    ) -> Path:
        """Synchronous DOCX generation for resumes."""
        from docx import Document
        from docx.shared import Pt

        try:
            doc = Document()
            self._set_default_font(doc, "Calibri", Pt(11))

            # Header: name and contact info
            self._add_header(doc, context)

            # Summary / Professional Summary
            summary = context.get("summary", "")
            if summary:
                doc.add_heading("Professional Summary", level=1)
                doc.add_paragraph(summary)

            # Experience
            self._add_experience(doc, context.get("experience", []))

            # Education
            self._add_education(doc, context.get("education", []))

            # Skills
            skills: list[str] = context.get("skills", [])
            if skills:
                doc.add_heading("Skills", level=1)
                doc.add_paragraph(", ".join(skills))

            # Certifications (optional)
            certs: list[str] = context.get("certifications", [])
            if certs:
                doc.add_heading("Certifications", level=1)
                for cert in certs:
                    doc.add_paragraph(cert, style="List Bullet")

            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

        except Exception as exc:
            raise GenerationError(
                f"DOCX rendering failed for template '{template_name}': {exc}",
            ) from exc

        logger.info(
            "docx_rendered",
            template=template_name,
            output=str(output_path),
            size_bytes=output_path.stat().st_size,
        )
        return output_path

    def _render_cover_letter_sync(
        self,
        content: str,
        output_path: Path,
    ) -> Path:
        """Synchronous DOCX generation for cover letters."""
        from docx import Document
        from docx.shared import Pt

        try:
            doc = Document()
            self._set_default_font(doc, "Georgia", Pt(12))

            for paragraph in content.split("\n\n"):
                text = paragraph.strip()
                if text:
                    doc.add_paragraph(text)

            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

        except Exception as exc:
            raise GenerationError(
                f"DOCX cover letter rendering failed: {exc}",
            ) from exc

        logger.info(
            "docx_cover_letter_rendered",
            output=str(output_path),
            size_bytes=output_path.stat().st_size,
        )
        return output_path

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _set_default_font(doc: Any, font_name: str, font_size: Any) -> None:
        """Set the default document font."""
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = font_size

    @staticmethod
    def _add_header(doc: Any, context: dict[str, Any]) -> None:
        """Add name and contact line to the document header."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        name = context.get("name", "")
        if name:
            heading = doc.add_heading(name, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_parts: list[str] = []
        for key in ("email", "phone", "location", "linkedin"):
            value = context.get(key)
            if value:
                contact_parts.append(value)

        if contact_parts:
            p = doc.add_paragraph(" | ".join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def _add_experience(doc: Any, experiences: list[dict[str, Any]]) -> None:
        """Add the Experience section."""
        if not experiences:
            return

        doc.add_heading("Experience", level=1)
        for exp in experiences:
            title = exp.get("title", "")
            company = exp.get("company", "")
            title_line = f"{title} — {company}" if company else title

            p = doc.add_paragraph()
            run = p.add_run(title_line)
            run.bold = True

            duration = exp.get("duration", "")
            if duration:
                p.add_run(f"  ({duration})")

            description = exp.get("description", "")
            if description:
                for bullet in description.split("\n"):
                    bullet = bullet.strip()
                    if bullet:
                        doc.add_paragraph(bullet, style="List Bullet")

    @staticmethod
    def _add_education(doc: Any, education: list[dict[str, Any]]) -> None:
        """Add the Education section."""
        if not education:
            return

        doc.add_heading("Education", level=1)
        for edu in education:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            line = f"{degree} — {institution}" if institution else degree

            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True

            year = edu.get("year", "")
            if year:
                p.add_run(f"  ({year})")
